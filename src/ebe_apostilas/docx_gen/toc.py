"""
Sumário automático (Índice) nativo do Microsoft Word.

Insere um campo TOC real (``TOC \\o "1-3" \\h \\z \\u``) que o Word/LibreOffice
actualiza automaticamente com base nos estilos de título (Heading 1/2/3)
usados no documento — em vez de uma lista estática, garantindo um índice
sempre coerente com a paginação real do ficheiro final.
"""
from __future__ import annotations

from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from ebe_apostilas.docx_gen.styles import COR_CITACAO, COR_PRIMARIA, FONTE_TITULO


def inserir_indice_automatico(doc: DocumentObject) -> None:
    """Insere um campo de Sumário/Índice automático do Word."""
    titulo_p = doc.add_paragraph()
    r = titulo_p.add_run("ÍNDICE")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    aviso_p = doc.add_paragraph()
    r = aviso_p.add_run(
        "(Clique com o botão direito sobre o índice e seleccione "
        "\u201cAtualizar campo\u201d para sincronizar a paginação, caso o "
        "documento seja editado.)"
    )
    r.font.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = COR_CITACAO

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    texto_marcador = OxmlElement("w:t")
    texto_marcador.text = (
        "O índice será gerado automaticamente ao abrir este documento no "
        "Microsoft Word (ou ao actualizar campos)."
    )

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fld_char_begin)
    r_element.append(instr_text)
    r_element.append(fld_char_separate)
    r_element.append(texto_marcador)
    r_element.append(fld_char_end)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
