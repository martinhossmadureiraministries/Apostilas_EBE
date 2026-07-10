"""
Construtor da apostila DOCX final.

Recebe um ``ItemCurricular`` (metadados oficiais) e um
``ConteudoApostilaGerado`` (texto validado, produzido pelo Gemini) e produz
um documento .docx completo, editável, com padrão editorial profissional:
capa, marco filosófico, ficha técnica, índice automático, cabeçalho e
rodapé com paginação, desenvolvimento estruturado, quadros de destaque,
tabelas, exercícios em três blocos, estudo bíblico complementar, glossário,
bibliografia e anotações pessoais.

Nenhuma apostila reutiliza texto de outra: todo o conteúdo textual é
injectado a partir do ``ConteudoApostilaGerado``, que é único por
apostila.
"""
from __future__ import annotations

import logging
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ebe_apostilas.core.config import Settings
from ebe_apostilas.core.models import ConteudoApostilaGerado, ItemCurricular
from ebe_apostilas.docx_gen.capa import add_capa_apostila, add_marco_filosofico
from ebe_apostilas.docx_gen.styles import (
    COR_SECUNDARIA,
    FONTE_TITULO,
    add_horizontal_line,
    citacao,
    h1,
    h2,
    h3,
    lista,
    novo_documento,
    page_break,
    paragrafo,
    quadro_destaque,
    selo_final,
    tabela_com_cabecalho,
)
from ebe_apostilas.docx_gen.toc import inserir_indice_automatico

logger = logging.getLogger("ebe_apostilas.docx_builder")


def _slugificar(texto: str, tamanho_maximo: int = 60) -> str:
    """Gera um nome de ficheiro seguro (sem acentos nem caracteres
    especiais) a partir de um título."""
    import unicodedata

    normalizado = unicodedata.normalize("NFKD", texto)
    sem_acentos = "".join(c for c in normalizado if not unicodedata.combining(c))
    limpo = "".join(c if c.isalnum() else "_" for c in sem_acentos)
    while "__" in limpo:
        limpo = limpo.replace("__", "_")
    return limpo.strip("_")[:tamanho_maximo]


class ApostilaDocxBuilder:
    """Monta e grava em disco o documento DOCX final de uma apostila."""

    def __init__(self, settings: Settings):
        self._settings = settings
        self._logo_path = settings.assets_dir / "logo_ebe.png"

    def caminho_saida(self, item: ItemCurricular) -> Path:
        nome_arquivo = f"{item.codigo}_{_slugificar(item.titulo)}.docx"
        return self._settings.output_dir / nome_arquivo

    def construir(self, item: ItemCurricular, conteudo: ConteudoApostilaGerado) -> Path:
        """Constrói o documento DOCX completo e devolve o caminho do
        ficheiro gravado."""
        self._settings.output_dir.mkdir(parents=True, exist_ok=True)

        doc = novo_documento(f"Apostila {item.codigo} — {item.titulo}", item.codigo)

        # ====== CAPA ======
        add_capa_apostila(
            doc,
            logo_path=self._logo_path,
            supratitulo=item.instituto_nome,
            titulo=conteudo.titulo,
            subtitulo=conteudo.subtitulo,
            codigo=item.codigo,
            numero_apostila=item.id,
            nivel_nome=item.nivel_nome,
            identificacao_extra=[
                ("Escola", item.escola_nome),
                ("Curso", f"{item.curso_nome} ({item.carga_horaria_curso})"),
                ("Módulo", f"{item.modulo_numero} — {item.modulo_nome}"),
            ],
        )

        # ====== MARCO FILOSÓFICO ======
        add_marco_filosofico(doc)

        # ====== FICHA TÉCNICA ======
        h1(doc, "Ficha Técnica")
        paragrafo(
            doc,
            "Este material didáctico é propriedade intelectual da Escola "
            "Bíblica Epignósis (EBE), produzido para uso exclusivo no "
            "âmbito dos seus programas de formação. A sua reprodução, no "
            "todo ou em parte, depende de autorização institucional escrita.",
        )
        lista(doc, [
            f"Título da apostila: {conteudo.titulo}.",
            f"Nível formativo: {item.nivel_nome}.",
            f"Instituto: {item.instituto_nome}.",
            f"Escola: {item.escola_nome}.",
            f"Curso: {item.curso_nome} (carga horária: {item.carga_horaria_curso}).",
            f"Módulo: {item.modulo_numero} — {item.modulo_nome}.",
            "Autor / Docente: Direcção Pedagógica da Escola Bíblica Epignósis.",
            "Revisão pedagógica: Coordenação Acadêmica.",
            "Revisão doutrinária: Conselho Doutrinário.",
            "Versão bíblica de referência: Almeida Revista e Corrigida (ARC).",
            "Edição: 1.ª — 2026.",
            f"Código institucional: {item.codigo}.",
        ])
        citacao(
            doc,
            "Toda a Escritura é divinamente inspirada e proveitosa para "
            "ensinar, para redarguir, para corrigir, para instruir em "
            "justiça; para que o homem de Deus seja perfeito e "
            "perfeitamente instruído para toda a boa obra.",
            "2 Timóteo 3.16-17",
        )
        page_break(doc)

        # ====== ÍNDICE AUTOMÁTICO ======
        inserir_indice_automatico(doc)
        page_break(doc)

        # ====== APRESENTAÇÃO ======
        h1(doc, "Apresentação da Apostila")
        for paragrafo_texto in _dividir_paragrafos(conteudo.apresentacao):
            paragrafo(doc, paragrafo_texto)

        # ====== OBJECTIVOS ======
        h1(doc, "Objectivos de Aprendizagem")
        paragrafo(doc, "Ao concluir o estudo desta apostila, o(a) aluno(a) será capaz de:")
        lista(doc, conteudo.objectivos, ordenada=True)

        # ====== VERSÍCULO-CHAVE ======
        h1(doc, "Versículo-Chave")
        citacao(doc, conteudo.versiculo_chave_texto, conteudo.versiculo_chave_referencia)

        # ====== TEXTO-BASE ======
        h1(doc, "Texto-Base para Leitura")
        paragrafo(
            doc,
            "Antes de iniciar o estudo, leia atentamente, em sua Bíblia "
            "(Almeida Revista e Corrigida), a seguinte passagem:",
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(conteudo.texto_base_referencia)
        r.font.name = FONTE_TITULO
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = COR_SECUNDARIA
        page_break(doc)

        # ====== INTRODUÇÃO ======
        h1(doc, "Introdução", numero=1)
        for paragrafo_texto in _dividir_paragrafos(conteudo.introducao):
            paragrafo(doc, paragrafo_texto)

        # ====== DESENVOLVIMENTO ======
        h1(doc, "Desenvolvimento do Conteúdo", numero=2)
        for indice, secao in enumerate(conteudo.desenvolvimento, 1):
            h2(doc, secao.titulo, numero=f"2.{indice}")
            for paragrafo_texto in _dividir_paragrafos(secao.conteudo):
                paragrafo(doc, paragrafo_texto)

        quadro_destaque(doc, conteudo.quadro_destaque_titulo, conteudo.quadro_destaque_texto)
        page_break(doc)

        # ====== APLICAÇÃO PRÁTICA ======
        h1(doc, "Aplicação Prática", numero=3)
        paragrafo(
            doc,
            "Esta secção convida o(a) aluno(a) a transformar o conhecimento "
            "adquirido em prática concreta, nas diversas esferas da vida cristã:",
        )
        lista(doc, conteudo.aplicacao_pratica, ordenada=True)

        # ====== SÍNTESE ======
        h1(doc, "Síntese e Conclusão", numero=4)
        for paragrafo_texto in _dividir_paragrafos(conteudo.sintese):
            paragrafo(doc, paragrafo_texto)
        page_break(doc)

        # ====== EXERCÍCIOS ======
        h1(doc, "Exercícios de Revisão")
        paragrafo(
            doc,
            "Responda às questões a seguir com base no conteúdo desta "
            "apostila e na sua leitura bíblica.",
        )
        h3(doc, "I — Verifique a sua compreensão")
        lista(doc, conteudo.exercicios_compreensao, ordenada=True)
        h3(doc, "II — Reflexão pessoal")
        lista(doc, conteudo.exercicios_reflexao, ordenada=True)
        h3(doc, "III — Ministério e serviço")
        lista(doc, conteudo.exercicios_ministerio, ordenada=True)

        # ====== ESTUDO BÍBLICO COMPLEMENTAR ======
        h1(doc, f"Estudo Bíblico Complementar — {conteudo.estudo_biblico_titulo}")
        for paragrafo_texto in _dividir_paragrafos(conteudo.estudo_biblico_texto):
            paragrafo(doc, paragrafo_texto)
        lista(doc, conteudo.estudo_biblico_perguntas, ordenada=True)

        # ====== RESUMO FINAL ======
        h1(doc, "Resumo Final")
        paragrafo(doc, conteudo.resumo_final)
        page_break(doc)

        # ====== GLOSSÁRIO ======
        h1(doc, "Glossário")
        paragrafo(doc, "Definições breves dos termos-chave utilizados nesta apostila.")
        tabela_com_cabecalho(
            doc,
            ["Termo", "Definição"],
            [(t.termo, t.definicao) for t in conteudo.glossario],
        )

        # ====== BIBLIOGRAFIA ======
        h1(doc, "Bibliografia Recomendada")
        lista(doc, conteudo.bibliografia)

        # ====== ANOTAÇÕES PESSOAIS ======
        h1(doc, "Anotações Pessoais")
        for _ in range(12):
            p = doc.add_paragraph()
            add_horizontal_line(p, color="C8C8C8", size=4)

        selo_final(doc)

        caminho_saida = self.caminho_saida(item)
        doc.save(str(caminho_saida))
        logger.info("Documento DOCX gravado: %s", caminho_saida)
        return caminho_saida


def _dividir_paragrafos(texto: str) -> list[str]:
    """Divide um bloco de texto gerado pelo Gemini em parágrafos
    individuais, respeitando quebras de linha duplas ou simples."""
    if not texto:
        return []
    partes = [p.strip() for p in texto.replace("\r\n", "\n").split("\n\n")]
    partes = [p for p in partes if p]
    if len(partes) <= 1:
        partes = [p.strip() for p in texto.split("\n") if p.strip()]
    return partes or [texto.strip()]
