"""
Módulo de estilos institucionais para os documentos DOCX da Escola Bíblica
Epignósis (EBE).

Design: académico formal (pt-PT, ARC), harmonizado com a identidade visual
oficial (azul-marinho `#1B3A5C` + verde `#2E7D4F` + dourado `#C9A14B`).

Este módulo consolida e reutiliza a biblioteca de estilos já validada nos
documentos institucionais da EBE (capa, marco filosófico, cabeçalhos,
rodapés com paginação automática, tabelas, citações, glossário etc.),
tornando-a reutilizável por qualquer gerador de conteúdo (apostilas,
manuais, avaliações, apresentações, planos de aula).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import _Cell

# === Paleta institucional (harmonizada com o logotipo EBE) ===
COR_PRIMARIA = RGBColor(0x1B, 0x3A, 0x5C)
COR_SECUNDARIA = RGBColor(0x2E, 0x7D, 0x4F)
COR_TERCIARIA = RGBColor(0xC9, 0xA1, 0x4B)
COR_TEXTO = RGBColor(0x1A, 0x1A, 0x1A)
COR_LINHA = RGBColor(0xB8, 0xB8, 0xB8)
COR_CITACAO = RGBColor(0x55, 0x55, 0x55)

HEX_PRIMARIA = "1B3A5C"
HEX_SECUNDARIA = "2E7D4F"
HEX_TERCIARIA = "C9A14B"

FONTE_TITULO = "Garamond"
FONTE_CORPO = "Garamond"


def _set_cell_border(cell: _Cell, **kwargs) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), kwargs[edge].get("val", "single"))
            border.set(qn("w:sz"), str(kwargs[edge].get("sz", 4)))
            border.set(qn("w:color"), kwargs[edge].get("color", "B8B8B8"))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def shade_cell(cell: _Cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_line(paragraph, color: str = HEX_PRIMARIA, size: int = 8) -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_page(section, top: float = 2.5, bottom: float = 2.5, left: float = 3.0, right: float = 2.5) -> None:
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def page_break(doc: DocumentObject) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def configurar_estilos_base(doc: DocumentObject) -> None:
    """Aplica os estilos académicos institucionais padrão (pt-PT)."""
    style = doc.styles["Normal"]
    style.font.name = FONTE_CORPO
    style.font.size = Pt(12)
    style.font.color.rgb = COR_TEXTO
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.4
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        _set_page(section)

    # Harmoniza os estilos nativos de título (usados pelo índice automático
    # do Word) com a identidade visual institucional.
    tamanhos = {"Heading 1": (Pt(15), COR_PRIMARIA), "Heading 2": (Pt(13), COR_PRIMARIA), "Heading 3": (Pt(11.5), COR_SECUNDARIA)}
    for nome_estilo, (tamanho, cor) in tamanhos.items():
        try:
            estilo = doc.styles[nome_estilo]
        except KeyError:
            continue
        estilo.font.name = FONTE_TITULO
        estilo.font.size = tamanho
        estilo.font.bold = True
        estilo.font.color.rgb = cor
        estilo.font.italic = False


def inserir_logo(doc: DocumentObject, caminho: Path, largura_cm: float = 5.5, alinhamento: str = "center"):
    p = doc.add_paragraph()
    if alinhamento == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alinhamento == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run()
    if caminho.exists():
        r.add_picture(str(caminho), width=Cm(largura_cm))
    return p


def h1(doc: DocumentObject, texto: str, numero: str | int | None = None):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    txt = f"{numero}. {texto}" if numero is not None else texto
    r = p.add_run(txt.upper())
    r.font.name = FONTE_TITULO
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA
    add_horizontal_line(p, color=HEX_PRIMARIA, size=6)
    return p


def h2(doc: DocumentObject, texto: str, numero: str | int | None = None):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    txt = f"{numero}. {texto}" if numero is not None else texto
    r = p.add_run(txt)
    r.font.name = FONTE_TITULO
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA
    return p


def h3(doc: DocumentObject, texto: str):
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(texto)
    r.font.name = FONTE_TITULO
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = COR_SECUNDARIA
    return p


def paragrafo(doc: DocumentObject, texto: str, italic: bool = False, bold: bool = False, justify: bool = True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texto)
    r.font.name = FONTE_CORPO
    r.font.size = Pt(12)
    r.font.italic = italic
    r.font.bold = bold
    return p


def citacao(doc: DocumentObject, texto: str, referencia: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"“{texto}”")
    r.font.name = FONTE_CORPO
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = COR_CITACAO
    if referencia:
        r2 = p.add_run(f"  ({referencia}, ARC)")
        r2.font.name = FONTE_CORPO
        r2.font.size = Pt(10)
        r2.font.italic = True
        r2.font.color.rgb = COR_SECUNDARIA
    return p


def lista(doc: DocumentObject, itens: list[str], ordenada: bool = False) -> None:
    for i, item in enumerate(itens, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(2)
        marca = f"{i}. " if ordenada else "•  "
        r = p.add_run(marca)
        r.font.name = FONTE_TITULO
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = COR_SECUNDARIA
        r2 = p.add_run(item)
        r2.font.name = FONTE_CORPO
        r2.font.size = Pt(12)


def quadro_destaque(doc: DocumentObject, titulo: str, texto: str) -> None:
    """Tabela de uma célula, sombreada, usada como quadro de destaque."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, "E8F1EC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"✦ {titulo}:  ")
    r.font.bold = True
    r.font.color.rgb = COR_SECUNDARIA
    r.font.name = FONTE_TITULO
    r.font.size = Pt(11)
    r2 = p.add_run(texto)
    r2.font.name = FONTE_CORPO
    r2.font.size = Pt(11)
    r2.font.italic = True


def tabela_com_cabecalho(doc: DocumentObject, cabecalhos: list[str], linhas: list[tuple]) -> None:
    tbl = doc.add_table(rows=1, cols=len(cabecalhos))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, titulo in enumerate(cabecalhos):
        hdr[i].text = ""
        shade_cell(hdr[i], HEX_PRIMARIA)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(titulo)
        r.font.bold = True
        r.font.name = FONTE_TITULO
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for linha in linhas:
        row = tbl.add_row().cells
        for i, valor in enumerate(linha):
            row[i].text = str(valor)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.name = FONTE_CORPO
                    r.font.size = Pt(10)


def selo_final(doc: DocumentObject) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_horizontal_line(p, color=HEX_SECUNDARIA, size=4)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESCOLA BÍBLICA EPIGNÓSIS")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Conhecer a Deus. Viver a Palavra. Manifestar o Reino.")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Soli Deo Gloria")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = COR_CITACAO


def cabecalho_rodape(doc: DocumentObject, titulo_doc: str, codigo_doc: str) -> None:
    """Cabeçalho discreto com o título do documento + rodapé com código e
    paginação automática (campo PAGE nativo do Word)."""
    for section in doc.sections:
        section.different_first_page_header_footer = True

        header = section.header
        ph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ph.text = ""
        r = ph.add_run(f"Escola Bíblica Epignósis  ·  {titulo_doc}")
        r.font.name = FONTE_TITULO
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = COR_SECUNDARIA

        footer = section.footer
        pf = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pf.add_run(f"{codigo_doc}  ·  ")
        r.font.name = FONTE_CORPO
        r.font.size = Pt(9)
        r.font.color.rgb = COR_CITACAO

        run = pf.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)
        run.font.name = FONTE_CORPO
        run.font.size = Pt(9)
        run.font.color.rgb = COR_CITACAO


def novo_documento(titulo_doc: str, codigo_doc: str) -> DocumentObject:
    """Cria um ``Document`` padronizado, já com estilos institucionais e
    cabeçalho/rodapé aplicados."""
    doc = Document()
    configurar_estilos_base(doc)
    cabecalho_rodape(doc, titulo_doc, codigo_doc)
    return doc
