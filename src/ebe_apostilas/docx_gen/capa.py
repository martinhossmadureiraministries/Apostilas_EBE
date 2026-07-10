"""
Elementos de abertura institucional dos documentos EBE: capa académica e
página de marco filosófico, reutilizados por todos os geradores de
materiais (apostilas, manuais, avaliações, apresentações etc.).
"""
from __future__ import annotations

from pathlib import Path

from docx.document import Document as DocumentObject
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from ebe_apostilas.docx_gen.styles import (
    COR_CITACAO,
    COR_PRIMARIA,
    COR_SECUNDARIA,
    COR_TEXTO,
    FONTE_CORPO,
    FONTE_TITULO,
    HEX_SECUNDARIA,
    add_horizontal_line,
    inserir_logo,
    page_break,
    shade_cell,
)

MARCO_FILOSOFICO_TEXTO = (
    "Acreditamos que o verdadeiro conhecimento de Deus transforma a mente "
    "pela verdade das Escrituras, o coração pela acção do Espírito Santo e "
    "a vida pelo compromisso de viver e anunciar o Evangelho de Jesus Cristo."
)


def add_capa_apostila(
    doc: DocumentObject,
    logo_path: Path,
    supratitulo: str,
    titulo: str,
    subtitulo: str,
    codigo: str,
    numero_apostila: int,
    nivel_nome: str,
    ano: str = "2026",
    identificacao_extra: list[tuple[str, str]] | None = None,
) -> None:
    """Capa académica sóbria e institucional, com logotipo, quadro de
    identificação e código do documento."""
    doc.add_paragraph()
    inserir_logo(doc, logo_path, largura_cm=5.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Conhecer a Deus. Viver a Palavra. Manifestar o Reino.")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph()
    add_horizontal_line(p, color=HEX_SECUNDARIA, size=6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(supratitulo.upper())
    r.font.name = FONTE_TITULO
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = COR_SECUNDARIA

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"APOSTILA N.º {numero_apostila:04d}")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo)
    r.font.name = FONTE_TITULO
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = COR_PRIMARIA

    if subtitulo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitulo)
        r.font.name = FONTE_TITULO
        r.font.size = Pt(13)
        r.font.italic = True
        r.font.color.rgb = COR_TEXTO

    doc.add_paragraph()
    doc.add_paragraph()

    dados = [("Autor / Docente", "Direcção Pedagógica · Escola Bíblica Epignósis")]
    if identificacao_extra:
        dados.extend(identificacao_extra)
    dados.append(("Nível formativo", nivel_nome))
    dados.append(("Edição / Ano", f"1.ª edição — {ano}"))

    tbl = doc.add_table(rows=len(dados), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(dados):
        row = tbl.rows[i].cells
        row[0].text = k
        row[1].text = v
        shade_cell(row[0], "E8F1EC")
        for p in row[0].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = FONTE_TITULO
                r.font.size = Pt(10)
                r.font.color.rgb = COR_PRIMARIA
        for p in row[1].paragraphs:
            for r in p.runs:
                r.font.name = FONTE_CORPO
                r.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    add_horizontal_line(p, color=HEX_SECUNDARIA, size=4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Material didáctico oficial · Código {codigo} · {ano}")
    r.font.name = FONTE_CORPO
    r.font.size = Pt(9)
    r.font.color.rgb = COR_CITACAO

    page_break(doc)


def add_marco_filosofico(doc: DocumentObject) -> None:
    """Página do marco filosófico institucional, exibida após a capa."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MARCO FILOSÓFICO")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = COR_SECUNDARIA

    p = doc.add_paragraph()
    add_horizontal_line(p, color=HEX_SECUNDARIA, size=4)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(2)
    p.paragraph_format.right_indent = Cm(2)
    r = p.add_run(f"“{MARCO_FILOSOFICO_TEXTO}”")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = COR_PRIMARIA

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— Escola Bíblica Epignósis —")
    r.font.name = FONTE_TITULO
    r.font.size = Pt(10)
    r.font.color.rgb = COR_CITACAO

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "“Até que todos cheguemos à unidade da fé e ao pleno conhecimento "
        "(ἐπίγνωσις) do Filho de Deus, a homem perfeito, à medida da "
        "estatura completa de Cristo.”"
    )
    r.font.name = FONTE_CORPO
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = COR_CITACAO
    r.add_break()
    r2 = p.add_run("Efésios 4.13")
    r2.font.name = FONTE_CORPO
    r2.font.size = Pt(10)
    r2.font.italic = True
    r2.font.color.rgb = COR_CITACAO

    page_break(doc)
