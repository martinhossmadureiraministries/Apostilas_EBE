"""Testes do gerador de documentos DOCX das apostilas."""
from __future__ import annotations

from pathlib import Path

from docx import Document

from ebe_apostilas.core.config import Settings
from ebe_apostilas.core.curriculum import carregar_curriculo
from ebe_apostilas.core.models import ConteudoApostilaGerado
from ebe_apostilas.docx_gen.builder import ApostilaDocxBuilder

REPO_ROOT = Path(__file__).resolve().parents[1]
CURRICULUM_PATH = REPO_ROOT / "data" / "curriculo_apostilas.json"


def _obter_item_exemplo():
    curriculo = carregar_curriculo(CURRICULUM_PATH)
    return curriculo.obter(1)


def test_builder_gera_arquivo_docx_valido(settings_teste: Settings, conteudo_exemplo: ConteudoApostilaGerado):
    item = _obter_item_exemplo()
    builder = ApostilaDocxBuilder(settings_teste)
    caminho = builder.construir(item, conteudo_exemplo)

    assert caminho.exists()
    assert caminho.suffix == ".docx"
    assert caminho.stat().st_size > 10_000

    # O documento deve poder ser reaberto pela biblioteca python-docx sem
    # erros — garante que o ficheiro OOXML gerado é estruturalmente válido.
    doc = Document(str(caminho))
    assert len(doc.paragraphs) > 20


def test_builder_inclui_todos_os_headings_esperados(settings_teste: Settings, conteudo_exemplo: ConteudoApostilaGerado):
    item = _obter_item_exemplo()
    builder = ApostilaDocxBuilder(settings_teste)
    caminho = builder.construir(item, conteudo_exemplo)

    doc = Document(str(caminho))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

    textos_esperados = [
        "FICHA TÉCNICA", "APRESENTAÇÃO DA APOSTILA", "OBJECTIVOS DE APRENDIZAGEM",
        "VERSÍCULO-CHAVE", "TEXTO-BASE PARA LEITURA", "EXERCÍCIOS DE REVISÃO",
        "RESUMO FINAL", "GLOSSÁRIO", "BIBLIOGRAFIA RECOMENDADA", "ANOTAÇÕES PESSOAIS",
    ]
    for esperado in textos_esperados:
        assert any(esperado in h for h in headings), f"Heading ausente: {esperado}"


def test_builder_inclui_tabelas_glossario_e_quadro_destaque(settings_teste: Settings, conteudo_exemplo: ConteudoApostilaGerado):
    item = _obter_item_exemplo()
    builder = ApostilaDocxBuilder(settings_teste)
    caminho = builder.construir(item, conteudo_exemplo)

    doc = Document(str(caminho))
    # Pelo menos 2 tabelas: quadro de destaque + tabela do glossário
    # (mais a tabela de identificação da capa).
    assert len(doc.tables) >= 3


def test_builder_contem_pagebreaks(settings_teste: Settings, conteudo_exemplo: ConteudoApostilaGerado):
    item = _obter_item_exemplo()
    builder = ApostilaDocxBuilder(settings_teste)
    caminho = builder.construir(item, conteudo_exemplo)

    import zipfile

    with zipfile.ZipFile(caminho) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert xml.count('w:type="page"') >= 5


def test_builder_nome_arquivo_baseado_no_codigo_e_titulo(settings_teste: Settings):
    item = _obter_item_exemplo()
    builder = ApostilaDocxBuilder(settings_teste)
    caminho = builder.caminho_saida(item)
    assert caminho.name.startswith(item.codigo)
    assert caminho.suffix == ".docx"


def test_builder_gera_documentos_diferentes_para_itens_diferentes(settings_teste: Settings, conteudo_exemplo: ConteudoApostilaGerado):
    curriculo = carregar_curriculo(CURRICULUM_PATH)
    item1 = curriculo.obter(1)
    item2 = curriculo.obter(2)
    builder = ApostilaDocxBuilder(settings_teste)

    conteudo1 = conteudo_exemplo.model_copy(update={"titulo": item1.titulo})
    conteudo2 = conteudo_exemplo.model_copy(update={"titulo": item2.titulo})

    caminho1 = builder.construir(item1, conteudo1)
    caminho2 = builder.construir(item2, conteudo2)

    assert caminho1 != caminho2
    assert caminho1.exists() and caminho2.exists()
